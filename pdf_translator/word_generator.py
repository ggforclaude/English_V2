import io
import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt
from PIL import Image


class WordGenerator:
    PAGE_WIDTH_INCHES = 6.0

    def __init__(self):
        self.doc = Document()
        self._set_margins()

    def _set_margins(self):
        for section in self.doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(3.0)
            section.right_margin  = Cm(3.0)

    def create(self, pages: list, output_path: str):
        out = Path(output_path).resolve()

        # 출력 경로 검증
        if out.suffix.lower() != ".docx":
            raise ValueError(f"출력 파일은 .docx 형식이어야 합니다: {out.suffix}")
        parent = out.parent
        if not parent.exists():
            parent.mkdir(parents=True, exist_ok=True)
        if not os.access(parent, os.W_OK):
            raise PermissionError(f"쓰기 권한이 없습니다: {parent}")

        for idx, page in enumerate(pages):
            if idx > 0:
                self.doc.add_page_break()
            self._add_page(page)

        self.doc.save(str(out))

    def _add_page(self, page: dict):
        elements = []

        for block in page.get("text_blocks", []):
            y = block["bbox"][1] if block.get("bbox") else 0
            elements.append(("text", y, block))

        imgs = page.get("processed_images") or page.get("images", [])
        for img in imgs:
            bbox = img.get("bbox")
            if bbox is not None:
                y = bbox.y0 if hasattr(bbox, "y0") else bbox[1]
            else:
                y = float("inf")
            elements.append(("image", y, img))

        elements.sort(key=lambda e: e[1])

        for kind, _, data in elements:
            if kind == "text":
                self._add_text(data)
            else:
                self._add_image(data)

    def _add_text(self, block: dict):
        text = (block.get("translated") or block.get("text", "")).strip()
        if not text:
            return

        font_size = block.get("font_size", 11)
        is_bold   = block.get("is_bold", False)

        if font_size >= 18:
            self.doc.add_heading(text, level=1)
        elif font_size >= 14:
            self.doc.add_heading(text, level=2)
        else:
            para = self.doc.add_paragraph()
            run  = para.add_run(text)
            run.bold = is_bold
            run.font.size = Pt(max(8, font_size))

    def _add_image(self, img_data: dict):
        raw = img_data.get("processed_bytes") or img_data.get("bytes")
        if not raw:
            return

        try:
            pil_img = Image.open(io.BytesIO(raw))
            iw, ih  = pil_img.size
            if iw <= 0 or ih <= 0:
                return

            width_in  = min(self.PAGE_WIDTH_INCHES, iw / 96)
            height_in = width_in * ih / iw

            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(io.BytesIO(raw), width=Inches(width_in))

        except Exception as e:
            print(f"  [이미지 삽입 경고] {e}")
